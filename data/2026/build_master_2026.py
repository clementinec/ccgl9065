#!/usr/bin/env python3
"""Build the 2026 portfolio master sheet from Moodle exports and extracts."""

from __future__ import annotations

import html
import re
from pathlib import Path

import pandas as pd
from docx import Document
from PyPDF2 import PdfReader


SCRIPT_PATH = Path(__file__).resolve()
DATA_DIR = SCRIPT_PATH.parent
REPO_ROOT = DATA_DIR.parents[1]
EXTRACT_DIR = DATA_DIR / "_extracted"

ASSIGNMENTS = {
    "essay": DATA_DIR / "180990675_Final_Portfolio_Submission_Essay_pdf_Moodle_TT_.xls",
    "collage": DATA_DIR / "180990674_Final_Portfolio_Submission_Collage_txt_Moodle_TT_.xls",
    "video": DATA_DIR / "185096096_Final_Portfolio_Submission_Video_txt_Moodle_TT_.xls",
}

EXTRACTS = {
    "essay": EXTRACT_DIR / "essay_extract.csv",
    "collage": EXTRACT_DIR / "collage_extract.csv",
    "video": EXTRACT_DIR / "video_extract.csv",
}

MASTER_XLSX = DATA_DIR / "portfolio_master_2026.xlsx"
MASTER_CSV = DATA_DIR / "portfolio_master_2026.csv"
ISSUES_CSV = DATA_DIR / "portfolio_master_2026_issues.csv"
NOTES_MD = DATA_DIR / "portfolio_master_2026_notes.md"


YOUTUBE_PATTERNS = [
    re.compile(r"youtube\.com/(?:embed/|shorts/)([A-Za-z0-9_-]{11})"),
    re.compile(r"youtube\.com/watch\?v=([A-Za-z0-9_-]{11})"),
    re.compile(r"youtu\.be/([A-Za-z0-9_-]{11})"),
    re.compile(r"youtube\.comwatchv[_=]?([A-Za-z0-9_-]{11})"),
    re.compile(r"youtu\.be([A-Za-z0-9_-]{11})"),
]

URL_RE = re.compile(r"https?://[^\s<>\"]+")

# These came in as ImgBB viewer-page links, but the public pages expose direct
# image URLs. Keep them as explicit overrides so the build stays offline-safe.
IBB_PAGE_DIRECT_OVERRIDES = {
    "https://ibb.co/BKnrWtBw": "https://i.ibb.co/6JX19mFN/Your-super-sized-burgers-are-choking-coral-reefs.png",
    "https://ibb.co/997wpZ0x": "https://i.ibb.co/kgt3GmTL/u3658380-CCGL9065-collage.png",
}


def rel(path: Path | str | float | None) -> str:
    if path is None or pd.isna(path) or str(path) == "":
        return ""
    path_obj = Path(str(path))
    if not path_obj.is_absolute():
        path_obj = REPO_ROOT / path_obj
    try:
        return path_obj.resolve().relative_to(REPO_ROOT).as_posix()
    except ValueError:
        return str(path)


def clean(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    return str(value).strip()


def normalize_name(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def full_name(row: pd.Series) -> str:
    return " ".join(part for part in [clean(row.get("First Name")), clean(row.get("Last Name"))] if part)


def parse_paper_id(source_path: str) -> str:
    return Path(source_path).name.split(" - ", 1)[0].strip()


def youtube_ids(text: str) -> list[str]:
    decoded = html.unescape(text or "")
    ids: list[str] = []
    for pattern in YOUTUBE_PATTERNS:
        for match in pattern.findall(decoded):
            if match not in ids:
                ids.append(match)
    return ids


def youtube_urls(text: str) -> list[str]:
    decoded = html.unescape(text or "")
    urls: list[str] = []
    for url in URL_RE.findall(decoded):
        url = url.rstrip(".,;)}]")
        if ("youtube.com" in url or "youtu.be" in url) and url not in urls:
            urls.append(url)
    return urls


def embed_url(video_id: str) -> str:
    return f"https://www.youtube.com/embed/{video_id}"


def infer_malformed_ibb(source_path: str) -> str:
    """Recover Moodle filenames like httpsi.ibb.copvRcd4zMCCGL9065-Collage.png."""
    name = Path(source_path).name
    match = re.search(r"httpsi\.ibb\.co([A-Za-z0-9]{8})([^_]+\.(?:png|jpe?g|gif))", name, re.I)
    if not match:
        return ""
    token, filename = match.groups()
    return f"https://i.ibb.co/{token}/{filename}"


def read_link_text(source_path: str) -> str:
    path = REPO_ROOT / source_path
    text = path.name
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".rtf", ".html", ".htm"}:
            return text + "\n" + path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            doc = Document(path)
            parts = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return text + "\n" + "\n".join(parts)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            parts = [page.extract_text() or "" for page in reader.pages]
            return text + "\n" + "\n".join(parts)
    except Exception as exc:  # Link scraping is best-effort; preserve the build.
        return text + f"\n[link text extraction failed: {exc}]"
    return text


def ibb_links(text: str) -> tuple[list[str], list[str]]:
    decoded = html.unescape(text or "")
    direct: list[str] = []
    page: list[str] = []
    for match in re.findall(r"https?://i\.ibb\.co/[^\s<>\"']+\.(?:png|jpe?g|gif)", decoded, re.I):
        cleaned = match.rstrip(".,;)}]")
        if cleaned not in direct:
            direct.append(cleaned)
    for match in re.findall(r"https?://ibb\.co/[A-Za-z0-9]+", decoded, re.I):
        cleaned = match.rstrip(".,;)}]")
        if cleaned not in page:
            page.append(cleaned)
    return direct, page


def add_note(current: str, note: str) -> str:
    if not note:
        return current
    if not current:
        return note
    if note in current:
        return current
    return f"{current} | {note}"


def empty_student() -> dict[str, str]:
    return {
        "student_name": "",
        "last_name": "",
        "first_name": "",
        "email": "",
        "student_number": "",
        "turnitin_user_id": "",
        "essay_paper_id": "",
        "essay_title": "",
        "essay_date_uploaded": "",
        "essay_pdf_path": "",
        "essay_source_path": "",
        "essay_source_ext": "",
        "essay_converted_from_docx": "",
        "essay_conversion_needed": "",
        "essay_notes": "",
        "collage_paper_id": "",
        "collage_title": "",
        "collage_date_uploaded": "",
        "collage_best_link_or_path": "",
        "collage_ibb_image_link": "",
        "collage_ibb_page_link": "",
        "collage_other_link": "",
        "local_collage_path": "",
        "collage_source_path": "",
        "collage_source_ext": "",
        "collage_cross_links_notes": "",
        "collage_notes": "",
        "video_paper_id": "",
        "video_title": "",
        "video_date_uploaded": "",
        "video_best_link_or_path": "",
        "youtube_embed_link": "",
        "raw_youtube_link": "",
        "local_video_path": "",
        "video_source_path": "",
        "video_source_ext": "",
        "video_source_kind": "",
        "video_notes": "",
    }


def assignment_prefix(assignment: str, column: str) -> str:
    return f"{assignment}_{column}"


def main() -> None:
    EXTRACT_DIR.mkdir(exist_ok=True)

    rosters: dict[str, pd.DataFrame] = {}
    paper_lookup: dict[tuple[str, str], pd.Series] = {}
    students: dict[str, dict[str, str]] = {}
    name_to_key: dict[str, str] = {}

    def key_for(row: pd.Series | None, extracted_name: str = "") -> str:
        email = clean(row.get("Email") if row is not None else "")
        name = full_name(row) if row is not None else extracted_name
        if email:
            return email.lower()
        return f"name:{normalize_name(name)}"

    def get_student(key: str, row: pd.Series | None = None, extracted_name: str = "") -> dict[str, str]:
        if key not in students:
            students[key] = empty_student()
        record = students[key]
        if row is not None:
            record["student_name"] = record["student_name"] or full_name(row)
            record["last_name"] = record["last_name"] or clean(row.get("Last Name"))
            record["first_name"] = record["first_name"] or clean(row.get("First Name"))
            record["email"] = record["email"] or clean(row.get("Email"))
            record["turnitin_user_id"] = record["turnitin_user_id"] or clean(row.get("Turnitin User ID"))
        elif extracted_name:
            record["student_name"] = record["student_name"] or extracted_name
        if record["student_name"]:
            name_to_key.setdefault(normalize_name(record["student_name"]), key)
        return record

    for assignment, path in ASSIGNMENTS.items():
        df = pd.read_excel(path)
        df["Paper ID"] = df["Paper ID"].astype(str)
        rosters[assignment] = df
        for _, row in df.iterrows():
            paper_id = clean(row.get("Paper ID"))
            paper_lookup[(assignment, paper_id)] = row
            key = key_for(row)
            record = get_student(key, row)
            record[assignment_prefix(assignment, "paper_id")] = paper_id
            record[assignment_prefix(assignment, "title")] = clean(row.get("Title"))
            record[assignment_prefix(assignment, "date_uploaded")] = clean(row.get("Date Uploaded"))

    for assignment, extract_path in EXTRACTS.items():
        if not extract_path.exists():
            continue
        extract = pd.read_csv(extract_path, dtype=str).fillna("")
        for _, row in extract.iterrows():
            paper_id = parse_paper_id(row["source_path"])
            roster_row = paper_lookup.get((assignment, paper_id))
            extracted_name = clean(row.get("student_name"))
            key = key_for(roster_row, extracted_name)
            if roster_row is None:
                key = name_to_key.get(normalize_name(extracted_name), key)
            record = get_student(key, roster_row, extracted_name)
            if clean(row.get("student_number")):
                record["student_number"] = clean(row.get("student_number"))

            source_path = rel(row.get("source_path"))
            source_ext = clean(row.get("source_ext")).lower()

            if assignment == "essay":
                record["essay_source_path"] = source_path
                record["essay_source_ext"] = source_ext
                record["essay_notes"] = add_note(record["essay_notes"], clean(row.get("notes")))
                if source_ext == "pdf":
                    record["essay_pdf_path"] = source_path
                    record["essay_conversion_needed"] = "false"
                else:
                    source_abs = REPO_ROOT / source_path
                    converted_abs = source_abs.with_suffix(".pdf")
                    if converted_abs.exists():
                        record["essay_pdf_path"] = rel(converted_abs)
                        record["essay_converted_from_docx"] = "true"
                        record["essay_conversion_needed"] = "false"
                        record["essay_notes"] = add_note(
                            record["essay_notes"], "DOCX converted to PDF beside original."
                        )
                    else:
                        record["essay_conversion_needed"] = "true"

            elif assignment == "collage":
                record["collage_source_path"] = source_path
                record["collage_source_ext"] = source_ext
                record["local_collage_path"] = rel(row.get("local_collage_path"))
                record["collage_cross_links_notes"] = clean(row.get("cross_links_notes"))
                record["collage_notes"] = add_note(record["collage_notes"], clean(row.get("notes")))

                candidate = clean(row.get("ibb_image_link"))
                inferred = infer_malformed_ibb(source_path)
                if candidate.startswith("https://i.ibb.co/"):
                    record["collage_ibb_image_link"] = candidate
                elif candidate.startswith("https://ibb.co/"):
                    record["collage_ibb_page_link"] = candidate
                    if candidate in IBB_PAGE_DIRECT_OVERRIDES:
                        record["collage_ibb_image_link"] = IBB_PAGE_DIRECT_OVERRIDES[candidate]
                        record["collage_notes"] = add_note(
                            record["collage_notes"], "Direct i.ibb.co URL recovered from public ibb.co page."
                        )
                elif candidate and ("canva.com" in candidate or "drive.google.com" in candidate):
                    record["collage_other_link"] = candidate
                elif candidate:
                    record["collage_notes"] = add_note(
                        record["collage_notes"], f"ignored non-collage URL candidate: {candidate}"
                    )

                if inferred and not record["collage_ibb_image_link"]:
                    record["collage_ibb_image_link"] = inferred
                    record["collage_notes"] = add_note(
                        record["collage_notes"], "i.ibb.co link inferred from malformed filename."
                    )

                if not record["collage_best_link_or_path"]:
                    record["collage_best_link_or_path"] = (
                        record["collage_ibb_image_link"]
                        or record["collage_ibb_page_link"]
                        or record["collage_other_link"]
                        or record["local_collage_path"]
                    )

                ids = youtube_ids(record["collage_cross_links_notes"])
                if ids and not record["youtube_embed_link"] and not record["local_video_path"]:
                    urls = youtube_urls(record["collage_cross_links_notes"])
                    record["youtube_embed_link"] = embed_url(ids[0])
                    record["raw_youtube_link"] = "; ".join(urls) if urls else embed_url(ids[0])
                    record["video_source_path"] = source_path
                    record["video_source_ext"] = source_ext
                    record["video_source_kind"] = "collage_cross_link"
                    record["video_notes"] = add_note(
                        record["video_notes"], "Video link recovered from collage submission."
                    )

            elif assignment == "video":
                record["video_source_path"] = source_path
                record["video_source_ext"] = source_ext
                record["video_source_kind"] = "video_submission"
                record["local_video_path"] = rel(row.get("local_video_path"))
                record["video_notes"] = add_note(record["video_notes"], clean(row.get("notes")))

                raw = clean(row.get("raw_youtube_link"))
                embed = clean(row.get("youtube_embed_link"))
                ids = youtube_ids(" ".join([embed, raw, source_path]))
                if not embed and ids:
                    embed = embed_url(ids[0])
                    raw = raw or f"https://www.youtube.com/watch?v={ids[0]}"
                    record["video_notes"] = add_note(
                        record["video_notes"], "YouTube link inferred from filename/source text."
                    )
                record["youtube_embed_link"] = embed
                record["raw_youtube_link"] = raw

                direct_ibb, page_ibb = ibb_links(read_link_text(source_path))
                if direct_ibb and not record["collage_ibb_image_link"]:
                    record["collage_ibb_image_link"] = direct_ibb[0]
                    record["collage_notes"] = add_note(
                        record["collage_notes"], "Collage image link recovered from video submission."
                    )
                elif page_ibb and not record["collage_ibb_page_link"]:
                    record["collage_ibb_page_link"] = page_ibb[0]
                    record["collage_notes"] = add_note(
                        record["collage_notes"], "Collage page link recovered from video submission."
                    )

    for record in students.values():
        record["collage_best_link_or_path"] = (
            record["collage_ibb_image_link"]
            or record["collage_ibb_page_link"]
            or record["collage_other_link"]
            or record["local_collage_path"]
        )
        record["video_best_link_or_path"] = record["youtube_embed_link"] or record["local_video_path"]

    master_columns = [
        "student_name",
        "last_name",
        "first_name",
        "email",
        "student_number",
        "turnitin_user_id",
        "collage_best_link_or_path",
        "collage_ibb_image_link",
        "collage_ibb_page_link",
        "collage_other_link",
        "local_collage_path",
        "essay_pdf_path",
        "essay_converted_from_docx",
        "essay_conversion_needed",
        "youtube_embed_link",
        "raw_youtube_link",
        "local_video_path",
        "video_best_link_or_path",
        "essay_source_path",
        "collage_source_path",
        "video_source_path",
        "video_source_kind",
        "essay_paper_id",
        "collage_paper_id",
        "video_paper_id",
        "essay_title",
        "collage_title",
        "video_title",
        "essay_date_uploaded",
        "collage_date_uploaded",
        "video_date_uploaded",
        "essay_notes",
        "collage_notes",
        "collage_cross_links_notes",
        "video_notes",
        "essay_source_ext",
        "collage_source_ext",
        "video_source_ext",
    ]

    master = pd.DataFrame(students.values()).fillna("")
    master = master[master_columns].sort_values(["last_name", "first_name", "student_name"])

    issue_rows: list[dict[str, str]] = []
    for _, record in master.iterrows():
        student = record["student_name"]
        if not record["essay_pdf_path"]:
            issue_rows.append({"student_name": student, "issue": "missing essay PDF", "detail": ""})
        if record["essay_conversion_needed"] == "true":
            issue_rows.append({"student_name": student, "issue": "essay DOCX still needs conversion", "detail": record["essay_source_path"]})
        if record["essay_converted_from_docx"] == "true":
            issue_rows.append({"student_name": student, "issue": "essay converted from DOCX", "detail": record["essay_pdf_path"]})
        if not record["collage_best_link_or_path"]:
            issue_rows.append({"student_name": student, "issue": "missing collage link/path", "detail": ""})
        elif not record["collage_ibb_image_link"]:
            issue_rows.append(
                {
                    "student_name": student,
                    "issue": "collage is not a direct i.ibb.co image link",
                    "detail": record["collage_best_link_or_path"],
                }
            )
        if not record["video_best_link_or_path"]:
            issue_rows.append({"student_name": student, "issue": "missing video embed/path", "detail": ""})
        elif not record["youtube_embed_link"]:
            issue_rows.append({"student_name": student, "issue": "video is a local upload without YouTube embed", "detail": record["local_video_path"]})
        if "ignored non-collage URL candidate" in record["collage_notes"]:
            issue_rows.append({"student_name": student, "issue": "ignored suspicious collage URL candidate", "detail": record["collage_notes"]})
        if "inferred from filename" in record["video_notes"] or "inferred from malformed filename" in record["collage_notes"]:
            issue_rows.append({"student_name": student, "issue": "link inferred from malformed source", "detail": "see notes columns"})

    issues = pd.DataFrame(issue_rows, columns=["student_name", "issue", "detail"])

    master.to_csv(MASTER_CSV, index=False)
    issues.to_csv(ISSUES_CSV, index=False)

    with pd.ExcelWriter(MASTER_XLSX, engine="xlsxwriter") as writer:
        master.to_excel(writer, index=False, sheet_name="master")
        issues.to_excel(writer, index=False, sheet_name="issues")
        for assignment, df in rosters.items():
            df.to_excel(writer, index=False, sheet_name=f"{assignment}_roster")
        for assignment, path in EXTRACTS.items():
            if path.exists():
                pd.read_csv(path, dtype=str).fillna("").to_excel(
                    writer, index=False, sheet_name=f"{assignment}_extract"
                )

        workbook = writer.book
        wrap = workbook.add_format({"text_wrap": True, "valign": "top"})
        for sheet_name, worksheet in writer.sheets.items():
            worksheet.freeze_panes(1, 0)
            worksheet.autofilter(0, 0, worksheet.dim_rowmax, worksheet.dim_colmax)
            worksheet.set_column(0, worksheet.dim_colmax, 18, wrap)
        writer.sheets["master"].set_column(0, 0, 30, wrap)
        writer.sheets["master"].set_column(6, 20, 48, wrap)
        writer.sheets["master"].set_column(31, 34, 55, wrap)
        writer.sheets["issues"].set_column(0, 0, 30, wrap)
        writer.sheets["issues"].set_column(1, 2, 60, wrap)

    note_lines = [
        "# 2026 portfolio master notes",
        "",
        f"- Master workbook: `{rel(MASTER_XLSX)}`",
        f"- Master CSV: `{rel(MASTER_CSV)}`",
        f"- Issues CSV: `{rel(ISSUES_CSV)}`",
        f"- Student rows: {len(master)}",
        f"- Rows with essay PDF paths: {(master['essay_pdf_path'] != '').sum()}",
        f"- Rows with direct i.ibb.co collage image links: {(master['collage_ibb_image_link'] != '').sum()}",
        f"- Rows with any collage link/path: {(master['collage_best_link_or_path'] != '').sum()}",
        f"- Rows with YouTube embed links: {(master['youtube_embed_link'] != '').sum()}",
        f"- Rows with local video paths: {(master['local_video_path'] != '').sum()}",
        f"- Issue rows: {len(issues)}",
        f"- Missing essay PDFs: {', '.join(master.loc[master['essay_pdf_path'] == '', 'student_name']) or 'none'}",
        f"- Missing collage link/path: {', '.join(master.loc[master['collage_best_link_or_path'] == '', 'student_name']) or 'none'}",
        f"- Missing video embed/path: {', '.join(master.loc[master['video_best_link_or_path'] == '', 'student_name']) or 'none'}",
        "",
        "The master sheet is keyed from the official Moodle workbooks when possible. "
        "Video links submitted in the collage folder are used as fallbacks when no video-folder submission exists.",
    ]
    NOTES_MD.write_text("\n".join(note_lines) + "\n", encoding="utf-8")

    print(f"Wrote {rel(MASTER_XLSX)}")
    print(f"Wrote {rel(MASTER_CSV)}")
    print(f"Wrote {rel(ISSUES_CSV)}")
    print(f"Wrote {rel(NOTES_MD)}")
    print(f"Student rows: {len(master)}; issue rows: {len(issues)}")


if __name__ == "__main__":
    main()
